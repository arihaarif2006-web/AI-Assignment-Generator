import os
import re
import threading
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googletrans import Translator
import tkinter as tk
from tkinter import messagebox

# ================================
# SETTINGS
# ================================

BASE_PATH = r"A:\ML_Project"
os.makedirs(BASE_PATH, exist_ok=True)

# ================================
# GOOGLE SEARCH FUNCTION
# ================================

def get_google_urls(query, num_results=10):
    urls = []
    headers = {"User-Agent": "Mozilla/5.0"}
    
    search_url = "https://html.duckduckgo.com/html/"
    
    try:
        response = requests.post(
            search_url,
            data={"q": query},
            headers=headers,
            timeout=10
        )
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        results = soup.find_all("a", class_="result__a")
        
        for result in results:
            url = result.get("href")
            
            if url and url.startswith("http"):
                if url not in urls:
                    urls.append(url)
            
            if len(urls) >= num_results:
                break
    
    except Exception as e:
        messagebox.showerror("Error", f"Search failed:\n{e}")
    
    return urls


# ================================
# SCRAPE CONTENT
# ================================

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def scrape_page(url):
    headers = {"User-Agent": "Mozilla/5.0"}
    text_data = ""
    
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        
        for tag in soup(["script", "style"]):
            tag.decompose()
        
        paragraphs = soup.find_all("p")
        
        for p in paragraphs:
            text = clean_text(p.get_text())
            if len(text) > 80:
                text_data += text + "\n"
                
    except:
        pass
    
    return text_data

# ================================
# GENERATE ASSIGNMENT
# ================================

def generate_assignment(topic, content):
    paragraphs = content.split("\n")
    chunk = max(1, len(paragraphs)//5)
    
    intro = f"{topic} is an important subject that plays a vital role in modern society. This assignment explores its major aspects and applications."
    
    headings = [
        "Introduction to the Topic",
        "Core Concepts",
        "Applications",
        "Advantages and Challenges",
        "Future Scope"
    ]
    
    sections = []
    for i in range(5):
        start = i * chunk
        end = start + chunk
        section_text = " ".join(paragraphs[start:end])
        sections.append((headings[i], section_text))
    
    conclusion = f"In conclusion, {topic} continues to grow in importance. Understanding its concepts and applications helps in future advancements."
    
    return intro, sections, conclusion

# ================================
# MAIN PROCESS FUNCTION
# ================================

def start_process():
    topic = topic_entry.get().strip()
    
    if not topic:
        messagebox.showwarning("Warning", "Please enter a topic.")
        return
    
    status_label.config(text="Processing... Please wait.")
    
    def run():
        try:
            # Step 1: Get URLs
            urls = get_google_urls(topic)
            
            if not urls:
                messagebox.showerror("Error", "No URLs found.")
                return
            
            # Save URLs
            with open(os.path.join(BASE_PATH, "urls.txt"), "w", encoding="utf-8") as f:
                for url in urls:
                    f.write(url + "\n")
            
            # Step 2: Scrape content
            combined_content = ""
            for url in urls:
                combined_content += scrape_page(url) + "\n"
            
            combined_content = combined_content[:20000]
            
            # Step 3: Generate English assignment
            intro, sections, conclusion = generate_assignment(topic, combined_content)
            
            doc = Document()
            title = doc.add_heading(topic, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading("Introduction", level=1)
            doc.add_paragraph(intro)
            
            for heading, text in sections:
                doc.add_heading(heading, level=1)
                doc.add_paragraph(text)
            
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(conclusion)
            
            english_path = os.path.join(BASE_PATH, f"{topic}_English.docx")
            doc.save(english_path)
            
            # Step 4: Translate to Urdu
            translator = Translator()
            urdu_doc = Document()
            
            for para in doc.paragraphs:
                if para.text.strip():
                    try:
                        translated = translator.translate(para.text, dest="ur").text
                        urdu_doc.add_paragraph(translated)
                    except:
                        urdu_doc.add_paragraph(para.text)
            
            urdu_path = os.path.join(BASE_PATH, f"{topic}_Urdu.docx")
            urdu_doc.save(urdu_path)
            
            status_label.config(text="Completed Successfully!")
            messagebox.showinfo("Success", "Assignment Generated Successfully!")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))
            status_label.config(text="Error occurred.")
    
    threading.Thread(target=run).start()

# ================================
# GUI DESIGN
# ================================

root = tk.Tk()
root.title("Automatic Assignment Generator")
root.geometry("500x250")
root.resizable(False, False)

title_label = tk.Label(root, text="Automatic Assignment Generator", font=("Arial", 16, "bold"))
title_label.pack(pady=15)

topic_label = tk.Label(root, text="Enter Assignment Topic:")
topic_label.pack()

topic_entry = tk.Entry(root, width=50)
topic_entry.pack(pady=5)

generate_button = tk.Button(root, text="Generate Assignment", command=start_process, bg="#4CAF50", fg="white")
generate_button.pack(pady=15)

status_label = tk.Label(root, text="")
status_label.pack()

root.mainloop()
